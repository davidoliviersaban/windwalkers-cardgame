#!/usr/bin/env python3

"""
Card Review Document Generator
Generates a Word document containing card images and comments.
"""

try:
    from docx import Document
    from docx.shared import Inches
except ImportError:
    print("Error: python-docx package is not installed.")
    print("Please run: pip install python-docx")
    exit(1)

import argparse
import os
from typing import List, Tuple

def get_images_and_comments(folder_path: str) -> List[Tuple[str, str]]:
    """Get all images and their comments from the specified folder."""
    images_and_comments = []
    if not os.path.exists(folder_path):
        raise FileNotFoundError(f"Folder not found: {folder_path}")
    
    files = []
    for filename in os.listdir(folder_path):
        if filename.lower().endswith(('.jpg', '.png', '.jpeg')):
            files.append(filename)

    files.sort()

    for filename in files:
        image_path = os.path.join(folder_path, filename)
        comment = f"Commentaire pour {filename}"
        images_and_comments.append((image_path, comment))
    
    return images_and_comments

def main():
    parser = argparse.ArgumentParser(
        description='Generate a Word document with card images and comments.'
    )
    parser.add_argument(
        'input_folder', 
        type=str, 
        help='Path to the folder containing card images'
    )
    parser.add_argument(
        '--output', 
        type=str, 
        default='card_review.docx',
        help='Output document name (default: card_review.docx)'
    )
    args = parser.parse_args()

    try:
        # Get images and comments
        images_and_comments = get_images_and_comments(args.input_folder)
        
        if not images_and_comments:
            print(f"No images found in {args.input_folder}")
            return

        # Create document
        doc = Document()
        doc.add_heading('Card Review Document', level=1)

        # Create table
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Add header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Card Image"
        header_cells[1].text = "Comments"

        # Add content
        for image_path, comment in images_and_comments:
            row_cells = table.add_row().cells
            paragraph = row_cells[0].paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(image_path, width=Inches(1.5))
            row_cells[1].text = comment

        # Save document
        doc.save(args.output)
        print(f"Document generated successfully: {args.output}")

    except Exception as e:
        print(f"Error: {str(e)}")
        exit(1)

if __name__ == "__main__":
    main()
